"""
Export Service — Exporters

TXT, DOCX, and PDF exporters.
PDF supports RTL Arabic text via NotoSansArabic font.
"""

from io import BytesIO
from fpdf import FPDF
from docx import Document
from shared.logger.logger import get_logger

logger = get_logger(__name__)


async def export_txt(content: str, title: str = "") -> bytes:
    """Export content as plain text."""
    output = f"{title}\n{'=' * len(title)}\n\n{content}" if title else content
    return output.encode("utf-8")


async def export_docx(content: str, title: str = "") -> bytes:
    """Export content as DOCX file."""
    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    for paragraph in content.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


async def export_pdf(content: str, title: str = "", language: str = "en") -> bytes:
    """
    Export content as PDF.
    Uses NotoSansArabic for RTL Arabic text.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # For Arabic, try to add Arabic font
    if language == "ar":
        try:
            pdf.add_font("NotoSansArabic", "", "fonts/NotoSansArabic-Regular.ttf", uni=True)
            pdf.set_font("NotoSansArabic", size=12)
        except Exception:
            logger.warning("arabic_font_not_found", fallback="Helvetica")
            pdf.set_font("Helvetica", size=12)
    else:
        pdf.set_font("Helvetica", size=12)

    if title:
        pdf.set_font_size(18)
        pdf.cell(0, 15, title, ln=True, align="C")
        pdf.set_font_size(12)
        pdf.ln(10)

    for line in content.split("\n"):
        pdf.multi_cell(0, 8, line)

    buffer = BytesIO()
    pdf.output(buffer)
    return buffer.getvalue()
